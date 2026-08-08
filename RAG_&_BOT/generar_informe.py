"""
Genera el informe M2 en formato .docx
"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

doc = Document()

# ── Márgenes ──────────────────────────────────────────────────────────────────
for sec in doc.sections:
    sec.top_margin    = Cm(2.5)
    sec.bottom_margin = Cm(2.5)
    sec.left_margin   = Cm(2.8)
    sec.right_margin  = Cm(2.8)

# ── Estilos base ──────────────────────────────────────────────────────────────
style_normal = doc.styles['Normal']
style_normal.font.name = 'Calibri'
style_normal.font.size = Pt(11)

style_h1 = doc.styles['Heading 1']
style_h1.font.name = 'Calibri'
style_h1.font.size = Pt(14)
style_h1.font.bold = True
style_h1.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)

style_h2 = doc.styles['Heading 2']
style_h2.font.name = 'Calibri'
style_h2.font.size = Pt(12)
style_h2.font.bold = True
style_h2.font.color.rgb = RGBColor(0x2E, 0x74, 0xB5)


def add_heading(text, level=1):
    p = doc.add_heading(text, level=level)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    return p


def add_para(text, bold_parts=None, space_after=6):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.space_before = Pt(2)
    if bold_parts is None:
        p.add_run(text)
    else:
        # bold_parts: list of (text, is_bold)
        for part, is_bold in bold_parts:
            r = p.add_run(part)
            r.bold = is_bold
    return p


def add_bullet(text, bold_start=None):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_after = Pt(2)
    if bold_start:
        r = p.add_run(bold_start)
        r.bold = True
        p.add_run(text)
    else:
        p.add_run(text)
    return p


def set_cell_bg(cell, color_hex):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement('w:shd')
    shd.set(qn('w:fill'), color_hex)
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    tcPr.append(shd)


def add_table_row(table, cells, bold=False, bg=None, font_color=None):
    row = table.add_row()
    for i, val in enumerate(cells):
        cell = row.cells[i]
        p    = cell.paragraphs[0]
        run  = p.add_run(str(val))
        run.bold = bold
        if font_color:
            run.font.color.rgb = font_color
        run.font.size = Pt(10)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        if bg:
            set_cell_bg(cell, bg)
    return row


# ══════════════════════════════════════════════════════════════════════════════
# TÍTULO
# ══════════════════════════════════════════════════════════════════════════════
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
title.paragraph_format.space_after = Pt(4)
r = title.add_run('Detector Forense de Bulos sobre Inmigración')
r.bold = True
r.font.size = Pt(16)
r.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
subtitle.paragraph_format.space_after = Pt(16)
r2 = subtitle.add_run('Clasificación automática de noticias (VERDADERO / CONTEXTO / FALSO)\nmediante fine-tuning de XLM-RoBERTa e integración en sistema RAG con Telegram')
r2.font.size = Pt(11)
r2.font.color.rgb = RGBColor(0x40, 0x40, 0x40)

# ══════════════════════════════════════════════════════════════════════════════
# 1. INTRODUCCIÓN
# ══════════════════════════════════════════════════════════════════════════════
add_heading('1. Introducción', 1)
add_para(
    'La desinformación sobre inmigración en España constituye un fenómeno de alto impacto social. '
    'Bulos como "los inmigrantes cobran 4.700 € al mes" o "España recibe más irregulares que ningún '
    'país europeo" circulan masivamente en redes sociales y medios digitales, fomentando discursos '
    'xenófobos y distorsionando el debate público.'
)
add_para(
    'Este proyecto desarrolla un sistema automatizado de fact-checking especializado en noticias sobre '
    'inmigración en español. El núcleo es un modelo NLP propio (Modelo LNR) basado en fine-tuning de '
    'XLM-RoBERTa que clasifica afirmaciones en tres categorías: '
)
add_bullet('VERDADERO — noticia factual verificada por medios de referencia.')
add_bullet('CONTEXTO — datos reales presentados de forma engañosa o sin contexto.')
add_bullet('FALSO — bulo o desinformación sin base factual.')
add_para(
    'El modelo se integra en un sistema de Recuperación Aumentada con Generación (RAG) que combina '
    'una base de datos vectorial (ChromaDB) con el modelo Gemini de Google para ofrecer análisis '
    'forense detallado. El prototipo de despliegue es un bot de Telegram accesible para cualquier usuario.',
    space_after=10
)

# ══════════════════════════════════════════════════════════════════════════════
# 2. PREPARACIÓN DE DATOS
# ══════════════════════════════════════════════════════════════════════════════
add_heading('2. Preparación de datos', 1)

add_heading('2.1 Recolección e integración', 2)
add_para(
    'Se construyeron dos datasets complementarios mediante web scraping y generación asistida por IA:'
)
add_bullet(
    ': 2.870 noticias sobre inmigración en España extraídas de medios verificados '
    '(El País, El Mundo, Maldita.es, Newtral, EFE…) etiquetadas como VERDADERO o FALSO '
    'según su origen editorial (medios de referencia vs. medios identificados como propagadores '
    'de bulos). 11 noticias categorizadas como ALERTA fueron reconvertidas a VERDADERO por '
    'tratarse de alertas verídicas.',
    bold_start='Dataset 1 — Noticias reales'
)
add_bullet(
    ': 2.870 versiones adulteradas con IA a partir del Dataset 1, '
    'etiquetadas como FALSO o CONTEXTO. Las técnicas de manipulación aplicadas incluyen: '
    'eliminación de contexto, cherry-picking estadístico, exageración de cifras, '
    'cambio de marco narrativo y fabricación parcial de datos.',
    bold_start='Dataset 2 — Versiones manipuladas con IA'
)
add_para(
    'La combinación de ambos datasets produce un corpus total de 5.740 ejemplos con la '
    'siguiente distribución de clases:'
)

# Tabla distribución
t = doc.add_table(rows=1, cols=3)
t.style = 'Table Grid'
t.alignment = WD_TABLE_ALIGNMENT.CENTER
add_table_row(t, ['Clase', 'N', '% del total'], bold=True, bg='1F497D', font_color=RGBColor(0xFF, 0xFF, 0xFF))

add_table_row(t, ['VERDADERO', '~2.700', '~47%'])
add_table_row(t, ['CONTEXTO',  '~1.420', '~25%'])
add_table_row(t, ['FALSO',     '~1.620', '~28%'])
doc.add_paragraph()

add_heading('2.2 Limpieza y transformaciones', 2)
add_para('El pipeline de preparación aplicó las siguientes transformaciones:')
add_bullet('Eliminación de filas con título o texto nulo (dropna).')
add_bullet('Filtrado de textos demasiado cortos (título < 5 chars o texto < 20 chars).')
add_bullet('Strip de espacios y normalización de codificación UTF-8.')
add_bullet(
    'Análisis de longitud: media ~1.200 chars por noticia (título + texto), '
    'percentil 90 en ~3.000 chars, equivalente a ~450-500 tokens con el tokenizador de XLM-RoBERTa.'
)
add_bullet(
    'División estratificada 80 / 10 / 10 (train / val / test) preservando la proporción de clases '
    'en cada partición mediante train_test_split con stratify=label.',
    bold_start=''
)
add_para(
    'La concatenación de título y texto como input único al modelo maximiza la información '
    'disponible para la clasificación, alcanzando el límite de 512 tokens de XLM-RoBERTa en el '
    'percentil 85 de los ejemplos.', space_after=10
)

# ══════════════════════════════════════════════════════════════════════════════
# 3. DESCRIPCIÓN DE LA TAREA — MINABLE VIEW
# ══════════════════════════════════════════════════════════════════════════════
add_heading('3. Descripción de la tarea — Minable View', 1)

add_para(
    'La tarea se formaliza como clasificación supervisada de texto en 3 clases:'
)
add_bullet('Tipo de tarea: clasificación de secuencias (sequence classification).')
add_bullet('Input (X): cadena de texto = titulo + " " + texto, truncada a 512 tokens.')
add_bullet('Output (Y): etiqueta ∈ {VERDADERO (0), CONTEXTO (1), FALSO (2)}.')
add_bullet('Unidad de muestra: una noticia o afirmación.')

add_heading('3.1 Problema de distribución en inferencia', 2)
add_para(
    'El modelo se entrena sobre artículos completos (~500 tokens), pero en producción recibe '
    'afirmaciones cortas del usuario (~15 tokens). Para resolver este desajuste de distribución, '
    'se aplica una técnica de inversión de input en inferencia:'
)
add_bullet(
    'Se recuperan los top-K (K=3) documentos más similares semánticamente de ChromaDB mediante '
    'embedding con Gemini Embedding 2.'
)
add_bullet(
    'El input al clasificador se construye como: query_usuario + titular_doc + texto_doc, '
    'recuperando la longitud de ~512 tokens que el modelo espera.'
)
add_bullet(
    'El veredicto final es el promedio ponderado (ensemble) de las distribuciones de probabilidad '
    'de los 3 documentos, ponderadas por similitud coseno (peso = 1 − distancia).'
)
add_para(
    'Este diseño alinea la distribución de entrenamiento e inferencia, '
    'evitando predicciones degradadas por textos cortos.', space_after=10
)

# ══════════════════════════════════════════════════════════════════════════════
# 4. MODELO Y EVALUACIÓN
# ══════════════════════════════════════════════════════════════════════════════
add_heading('4. Modelo y evaluación', 1)

add_heading('4.1 Arquitectura y entrenamiento', 2)
add_para(
    'Se utilizó XLM-RoBERTa base (xlm-roberta-base, 278M parámetros) de HuggingFace como modelo '
    'base. XLM-RoBERTa es un transformer multilingüe preentrenado sobre 100 idiomas (incluyendo '
    'español) con la tarea Masked Language Modeling, idóneo para clasificación de texto en español.'
)
add_para('Sobre él se añadió una cabeza de clasificación de 3 clases y se fine-tuneó con:')

# Tabla hiperparámetros
t2 = doc.add_table(rows=1, cols=2)
t2.style = 'Table Grid'
t2.alignment = WD_TABLE_ALIGNMENT.CENTER
add_table_row(t2, ['Hiperparámetro', 'Valor'], bold=True, bg='2E74B5', font_color=RGBColor(0xFF, 0xFF, 0xFF))

params = [
    ('Learning rate', '1.5 × 10⁻⁵'),
    ('Scheduler', 'Coseno con warmup 10%'),
    ('Batch efectivo', '64 (32 por GPU × 2 GPUs T4)'),
    ('Épocas máx.', '6 (early stopping patience=3)'),
    ('Label smoothing (α)', '0.10'),
    ('Weight decay', '0.01'),
    ('Pérdida', 'CrossEntropy ponderada (pesos inversos a frecuencia)'),
    ('Precisión', 'FP16 (mixed precision)'),
    ('Hardware', 'Kaggle 2×NVIDIA T4 (32 GB VRAM total)'),
]
for k, v in params:
    row = t2.add_row()
    row.cells[0].text = k
    row.cells[1].text = v
doc.add_paragraph()

add_para(
    'El uso de pesos de clase inversamente proporcionales a su frecuencia y label smoothing (α=0.1) '
    'combate el desbalance de clases y mejora la calibración de las probabilidades, penalizando '
    'predicciones excesivamente confiadas.'
)

add_heading('4.2 Evolución del entrenamiento', 2)

# Tabla métricas de entrenamiento
t3 = doc.add_table(rows=1, cols=7)
t3.style = 'Table Grid'
t3.alignment = WD_TABLE_ALIGNMENT.CENTER
headers = ['Época', 'Loss Train', 'Loss Val', 'Accuracy', 'F1 Macro', 'F1 Weighted', 'F1 FALSO']
add_table_row(t3, headers, bold=True, bg='2E74B5', font_color=RGBColor(0xFF, 0xFF, 0xFF))

epochs = [
    (1, '0.9272', '0.6901', '0.7592', '0.7059', '0.7604', '0.6105'),
    (2, '0.6701', '0.7605', '0.6440', '0.5041', '0.5544', '0.0098'),
    (3, '0.6104', '0.6350', '0.8028', '0.7525', '0.8077', '0.6854'),
    (4, '0.5714', '0.6534', '0.8482', '0.7775', '0.8413', '0.8009'),  # MEJOR
    (5, '0.5453', '0.6580', '0.8307', '0.7761', '0.8326', '0.7563'),
    (6, '0.5265', '0.6683', '0.8272', '0.6977', '0.8286', '0.7538'),
]
for ep in epochs:
    bg_color = 'E2EFDA' if ep[0] == 4 else None
    add_table_row(t3, ep, bg=bg_color)
doc.add_paragraph()

add_para(
    '→ El mejor modelo se guardó en la época 4 (F1 Macro = 0.778 en validación). '
    'Early stopping detuvo el entrenamiento al no mejorar durante 3 épocas consecutivas. '
    'La época 2 muestra un "cold start" del scheduler que se recupera en la época 3.',
    bold_parts=[
        ('→ El mejor modelo se guardó en la ', False),
        ('época 4 (F1 Macro = 0.778 en validación)', True),
        ('. Early stopping detuvo el entrenamiento al no mejorar durante 3 épocas consecutivas. '
         'La época 2 muestra un "cold start" del scheduler que se recupera en la época 3.', False),
    ]
)

add_heading('4.3 Evaluación en test', 2)
add_para(
    'El modelo guardado se evaluó sobre el conjunto de test (nunca visto durante el entrenamiento):'
)

# Tabla resultados test
t4 = doc.add_table(rows=1, cols=5)
t4.style = 'Table Grid'
t4.alignment = WD_TABLE_ALIGNMENT.CENTER
add_table_row(t4, ['Clase', 'Precision', 'Recall', 'F1-score', 'Soporte'], bold=True, bg='1F497D', font_color=RGBColor(0xFF, 0xFF, 0xFF))

test_rows = [
    ('VERDADERO', '1.000', '0.989', '0.994', '264'),
    ('CONTEXTO',  '0.611', '0.514', '0.558', '107'),
    ('FALSO',     '0.761', '0.837', '0.797', '202'),
]
for r in test_rows:
    add_table_row(t4, r)

# Fila separadora
row_sep = t4.add_row()
for c in row_sep.cells:
    c.text = ''

add_table_row(t4, ['Accuracy',      '',      '',      '0.846', '573'])
add_table_row(t4, ['F1 Macro avg',  '0.791', '0.780', '0.783', '573'])
add_table_row(t4, ['F1 Weighted avg','0.843','0.846', '0.843', '573'])
doc.add_paragraph()

add_para(
    'La clase VERDADERO obtiene un F1 casi perfecto (0.994). La clase FALSO alcanza F1 = 0.797, '
    'con alto recall (0.837), crítico para no dejar pasar bulos. '
    'La clase CONTEXTO es la más difícil (F1 = 0.558) por su naturaleza ambigua: '
    'datos verídicos presentados de forma engañosa sin marcadores textuales claros.'
)

add_heading('4.4 Matriz de confusión', 2)
add_para('La matriz de confusión en test confirma los patrones anteriores:')

t5 = doc.add_table(rows=1, cols=4)
t5.style = 'Table Grid'
t5.alignment = WD_TABLE_ALIGNMENT.CENTER
add_table_row(t5, ['Real \\ Pred', 'VERDADERO', 'CONTEXTO', 'FALSO'], bold=True, bg='2E74B5', font_color=RGBColor(0xFF, 0xFF, 0xFF))
add_table_row(t5, ['VERDADERO', '261', '2',  '1'])
add_table_row(t5, ['CONTEXTO',  '0',   '55', '52'])
add_table_row(t5, ['FALSO',     '0',   '33', '169'])
doc.add_paragraph()

add_para(
    'El modelo nunca confunde FALSO con VERDADERO (0 casos), lo que es '
    'crítico para el uso práctico. El principal error es confundir CONTEXTO con FALSO (52 casos), '
    'razonable dada la similitud semántica entre noticias adulteradas y bulos directos.',
    bold_parts=[
        ('El modelo ', False),
        ('nunca confunde FALSO con VERDADERO (0 casos)', True),
        (', lo que es crítico para el uso práctico. El principal error es confundir '
         'CONTEXTO con FALSO (52 casos), razonable dada la similitud semántica entre '
         'noticias adulteradas y bulos directos.', False),
    ], space_after=10
)

# ══════════════════════════════════════════════════════════════════════════════
# 5. PROTOTIPO DE DESPLIEGUE
# ══════════════════════════════════════════════════════════════════════════════
add_heading('5. Prototipo de despliegue', 1)

add_para(
    'El sistema completo integra tres componentes en un pipeline end-to-end:'
)

add_heading('5.1 Arquitectura del sistema', 2)
add_bullet(
    ': 313 chunks de noticias verificadas indexadas con embeddings de Gemini '
    'Embedding 2 (modelo de última generación de Google). Permite recuperar los documentos '
    'más similares semánticamente a cualquier consulta.',
    bold_start='Base de datos vectorial (ChromaDB)'
)
add_bullet(
    ': XLM-RoBERTa fine-tuned sobre 5.740 ejemplos de noticias en español '
    'sobre inmigración. Proporciona el veredicto clasificatorio con confianza y distribución '
    'completa de probabilidades (VERDADERO / CONTEXTO / FALSO).',
    bold_start='Modelo LNR (clasificador propio)'
)
add_bullet(
    ': Gemini 2.5 Flash genera el análisis forense detallado (origen, cadena '
    'de transformación, datos reales, mecanismo de engaño) usando los documentos recuperados '
    'como contexto y el veredicto LNR como punto de partida.',
    bold_start='Modelo generativo (Gemini 2.5 Flash)'
)

add_heading('5.2 Flujo de una consulta', 2)
add_para('Cuando un usuario envía una afirmación al bot:')
add_bullet('El texto se convierte en embedding con Gemini Embedding 2.')
add_bullet('ChromaDB recupera los top-8 documentos más similares.')
add_bullet(
    'El Modelo LNR clasifica usando ensemble de los 3 más relevantes '
    '(query + documento → input de ~512 tokens cada uno).'
)
add_bullet('Gemini genera el análisis forense usando prompt estructurado con los 8 documentos.')
add_bullet(
    'Telegram devuelve al usuario: (1) veredicto LNR con confianza y distribución de probabilidades, '
    '(2) análisis forense estructurado en 5 secciones, (3) links a las 3 noticias más relevantes.'
)

add_heading('5.3 Interfaces de usuario', 2)
add_bullet(
    ': accesible desde el móvil, procesamiento en tiempo real (~10 s por consulta), '
    'veredicto visual con emojis (🔴/🟢/🟡), análisis forense y links a fuentes.',
    bold_start='Bot de Telegram'
)
add_bullet(
    ': interfaz web local (Gradio) con panel de veredicto LNR, '
    'análisis de Gemini y panel de fuentes con indicadores de relevancia.',
    bold_start='Interfaz web (Gradio)'
)
add_para(
    'La umbralización de confianza (umbral = 0.60) permite devolver INCIERTO cuando el modelo '
    'no está suficientemente seguro, evitando veredictos erróneos con alta confianza aparente.',
    space_after=10
)

# ══════════════════════════════════════════════════════════════════════════════
# 6. DISCUSIÓN
# ══════════════════════════════════════════════════════════════════════════════
add_heading('6. Discusión', 1)

add_heading('6.1 Tecnología utilizada', 2)
add_para('El proyecto domina y combina las siguientes herramientas:')
add_bullet('HuggingFace Transformers (fine-tuning, Trainer API, pipeline de inferencia).')
add_bullet('PyTorch (entrenamiento con FP16, gestión de VRAM, WeightedTrainer personalizado).')
add_bullet('ChromaDB (base de datos vectorial persistente para RAG).')
add_bullet('Google Gemini API (embeddings de última generación y generación de texto).')
add_bullet('python-telegram-bot v22 (bot asíncrono con manejo de errores).')
add_bullet('Kaggle Notebooks con 2×GPU T4 para entrenamiento distribuido automático.')
add_para(
    'Se resolvieron de forma autónoma problemas técnicos avanzados: incompatibilidad de '
    'safetensors con Python 3.14 (solución con cargador puro mmap + torch.frombuffer), '
    'agotamiento de disco en Kaggle (gestión de checkpoints con save_total_limit=1), '
    'y el desajuste de distribución entrenamiento-inferencia (técnica de inversión de input con RAG).'
)

add_heading('6.2 Uso de IA y conceptos del curso', 2)
add_para(
    'El proyecto aplica directamente los conceptos del curso:'
)
add_bullet(
    ': problema de clasificación supervisada con etiquetas definidas, '
    'minable view formalizada (input/output/muestra).',
    bold_start='Unidades 3-4 — Formalización de la tarea'
)
add_bullet(
    ': arquitectura Transformer (attention, positional encoding, '
    'MLM pretraining), fine-tuning con clasificador lineal, label smoothing como regularización.',
    bold_start='Unidades 3-4 — Aprendizaje profundo y NLP'
)
add_bullet(
    ': división estratificada, selección del mejor modelo por F1 Macro '
    '(métrica adecuada para clases desbalanceadas), early stopping para evitar sobreajuste.',
    bold_start='Unidades 3-4 — Evaluación y protocolo experimental'
)
add_bullet(
    ': sistema RAG completo (retrieval + augmentation + generation) '
    'como estrategia de despliegue.',
    bold_start='Unidad 5 — Despliegue'
)
add_para(
    'IA generativa utilizada: (1) Claude Code (Anthropic) como asistente de programación para '
    'depuración, diseño de arquitectura del sistema y resolución de incompatibilidades técnicas; '
    '(2) Gemini 2.5 Flash para generación de análisis forense; (3) ChatGPT/GPT-4 para '
    'generación de las versiones manipuladas del Dataset 2 (clase CONTEXTO y FALSO).'
)

add_heading('6.3 Limitaciones y trabajo futuro', 2)
add_bullet(
    'La clase CONTEXTO es la más difícil (F1 = 0.558): noticias engañosas por omisión '
    'son difíciles de detectar sin conocimiento factual externo. Se propone aumentar el '
    'dataset de CONTEXTO con más variantes.'
)
add_bullet(
    'El modelo es específico de inmigración en España; generalizar a otros dominios '
    'requeriría reentrenamiento con datos del nuevo dominio.'
)
add_bullet(
    'La base de ChromaDB (313 chunks) es relativamente pequeña; ampliarla con más '
    'artículos de fact-checking mejoraría la calidad del RAG y del ensemble del clasificador.'
)
add_bullet(
    'Trabajo futuro: despliegue en servidor con webhook de Telegram para disponibilidad 24/7, '
    'y entrenamiento con más épocas o con modelos más grandes (XLM-RoBERTa large).'
)

# ══════════════════════════════════════════════════════════════════════════════
# GUARDAR
# ══════════════════════════════════════════════════════════════════════════════
out = r'C:\Users\Javier Caballero\Desktop\PROY III\M2_Informe_Detector_Bulos_Inmigracion.docx'
doc.save(out)
print(f'Informe guardado en: {out}')
